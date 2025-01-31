import os
import shutil
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import requests
from bs4 import BeautifulSoup
from docx import Document
from datetime import datetime
import time
import schedule

# Gmail
def send_email_with_attachment(sender, receiver, subject, body, filename):
    msg = MIMEMultipart()
    msg["Subject"] = subject
    msg["From"] = sender
    msg["To"] = receiver
    msg.attach(MIMEText(body, 'plain'))

    if os.path.exists(filename):
        with open(filename, "rb") as attachment:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f"attachment; filename={os.path.basename(filename)}")
            msg.attach(part)
    else:
        print(f"Error: No se encontró el archivo {filename}")
        return
    
    try:
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(sender, "fevj cqwa idzz xjsn")  # Usa una App Password
            server.sendmail(sender, receiver, msg.as_string())
        print("Correo enviado con éxito.")
    except Exception as e:
        print("Error al enviar correo:", e)

# Scrape sitios web
def web_scraping():
    url = "https://weather.com/es-ES/tiempo/hoy/l/0fa2ee94fcc6dedc7ce14ec2216673b74efe0d1bed6f6c089197dcbba3095b83"
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')

    try:
        temperature = soup.find('span', attrs={"dir": "ltr"}).text
        description = soup.find('div', attrs={"data-testid" : "wxPhrase"}).text
        humidity_elem = soup.find('span', attrs={"data-testid": "PercentageValue"})
        humidity = humidity_elem.get_text(strip=True) if humidity_elem else "No disponible"


        doc = Document()
        doc.add_heading('Pronóstico del tiempo', 0)
        doc.add_paragraph(f"Temperatura: {temperature}")
        doc.add_paragraph(f"Descripción: {description}")
        doc.add_paragraph(f"Humedad: {humidity}")
        doc.save("pronostico_tiempo.docx")

        print("Información del tiempo guardada en 'pronostico_tiempo.docx'")
    except AttributeError:
        print("Error al obtener datos del clima. Verifica si la estructura de la página ha cambiado.")

# Cambiar rutas
def backup_files():
    source_folder = "C:/Users/Alexa/Documents"
    backup_folder = "C:/Users/Alexa/Backup"
    max_backups = 2
    
    if not os.path.exists(backup_folder):
        os.makedirs(backup_folder)

    backup_subfolders = [f for f in os.listdir(backup_folder) if os.path.isdir(os.path.join(backup_folder, f))]
    backup_subfolders.sort(key=lambda x: os.path.getctime(os.path.join(backup_folder, x)))
    
    while len(backup_subfolders) >= max_backups:
        folder_to_delete = backup_subfolders.pop(0)
        shutil.rmtree(os.path.join(backup_folder, folder_to_delete))
        print(f"Se eliminó la copia de seguridad antigua: {folder_to_delete}")
    
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    new_backup_folder = os.path.join(backup_folder, f"Backup_{timestamp}")
    os.makedirs(new_backup_folder)
    
    for file_name in os.listdir(source_folder):
        full_file_name = os.path.join(source_folder, file_name)
        if os.path.isfile(full_file_name):
            shutil.copy(full_file_name, new_backup_folder)
    
    print(f"Copia de seguridad completada en: {new_backup_folder}")
#Cambiar rutas
def organize_files():
    folder_path = "C:/Users/Alexa/Downloads"
    extensions = {
        "Imágenes": [".jpg", ".png", ".gif"],
        "Documentos": [".pdf", ".docx", ".txt"],
        "Videos": [".mp4", ".avi"],
    }

    for category, ext_list in extensions.items():
        category_path = os.path.join(folder_path, category)
        if not os.path.exists(category_path):
            os.makedirs(category_path)

        for file in os.listdir(folder_path):
            file_ext = os.path.splitext(file)[1].lower()
            if file_ext in ext_list:
                shutil.move(os.path.join(folder_path, file), category_path)
    print("Archivos organizados.")

# Sin schedule
def main_without_schedule():
    current_date = datetime.now().strftime("%d de %B")
    subject = f"Pronóstico del tiempo - {current_date}"

    web_scraping()
    backup_files()
    organize_files()
    
    if os.path.exists("pronostico_tiempo.docx"):
        send_email_with_attachment("alexandrisman18@gmail.com", "alexandrisman18@gmail.com", subject, "Ejemplo de como se podria usar esto:\n con google calendar para ver lo que hoy toca\n automatizar tareas\n enviar a compañeros copias de seguridad diaramente...etc", "pronostico_tiempo.docx")
    else:
        print("No se pudo enviar el correo porque el archivo no fue generado.")

def run_with_schedule():
    current_date = datetime.now().strftime("%d %B")
    subject = f"Pronóstico del tiempo - {current_date}"

    schedule.every().day.at("08:00").do(web_scraping)
    schedule.every().day.at("20:00").do(backup_files)
    schedule.every().day.at("23:00").do(organize_files)
    schedule.every().day.at("09:00").do(lambda: send_email_with_attachment("alexandrisman18@gmail.com", "alexandrisman18@gmail.com", subject, "Lo que hoy toca", "pronostico_tiempo.docx"))

    while True:
        schedule.run_pending()
        time.sleep(60)

# Ejecución principal
main_without_schedule()
# Para ejecutar con schedule, van por horas. (Mas orientado a ordenadores 24/7 en ejecucion)
# run_with_schedule()
